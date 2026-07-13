"""Render a tailored CV's structured content as downloadable PDF/DOCX documents.

Renders directly from the structured entities (contact header + sections that
carry either prose ``body`` or a list of ``entries``) rather than converting
from an intermediate HTML form. Both libraries are pure-Python with no system
dependencies, so the backend needs no native packages at deploy time.
"""

from __future__ import annotations

import io
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.domain.entities import (
    TailoredCV,
    TailoredCVContact,
    TailoredCVEntry,
    TailoredCVSection,
)

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Slate/indigo palette — muted, print-friendly, ATS-safe (colour is decorative).
_INK = HexColor("#0f172a")
_TEXT = HexColor("#334155")
_MUTED = HexColor("#64748b")
_ACCENT = HexColor("#1e3a8a")
_RULE = HexColor("#cbd5e1")

_FONT = "Helvetica"
_FONT_BOLD = "Helvetica-Bold"
_FONT_ITALIC = "Helvetica-Oblique"


def _styles() -> dict[str, ParagraphStyle]:
    return {
        "name": ParagraphStyle(
            "Name", fontName=_FONT_BOLD, fontSize=19, leading=22,
            textColor=_INK, alignment=TA_CENTER, spaceAfter=4,
        ),
        "contact": ParagraphStyle(
            "Contact", fontName=_FONT, fontSize=9, leading=13,
            textColor=_MUTED, alignment=TA_CENTER, spaceAfter=2,
        ),
        "heading": ParagraphStyle(
            "Heading", fontName=_FONT_BOLD, fontSize=11, leading=14,
            textColor=_ACCENT, spaceBefore=12, spaceAfter=2,
        ),
        "body": ParagraphStyle(
            "Body", fontName=_FONT, fontSize=9.5, leading=14,
            textColor=_TEXT, alignment=TA_JUSTIFY, spaceAfter=5,
        ),
        "entry_title": ParagraphStyle(
            "EntryTitle", fontName=_FONT_BOLD, fontSize=10, leading=13,
            textColor=_INK, spaceBefore=6,
        ),
        "entry_date": ParagraphStyle(
            "EntryDate", fontName=_FONT, fontSize=9, leading=13,
            textColor=_MUTED, alignment=2,  # TA_RIGHT
        ),
        "context": ParagraphStyle(
            "Context", fontName=_FONT_ITALIC, fontSize=8.5, leading=11,
            textColor=_MUTED, spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "Bullet", fontName=_FONT, fontSize=9.5, leading=13.5,
            textColor=_TEXT, alignment=TA_JUSTIFY,
        ),
    }


def _esc(text: str) -> str:
    return _xml_escape(text or "")


def _context_text(raw: str) -> str:
    """Normalize an entry's context descriptor to a bare phrase.

    The model sometimes echoes the base resume's "[Context: ...]" literal into
    the field; strip any surrounding brackets and a leading "Context:" so the
    renderer can format it consistently as "[Context: ...]".
    """

    text = raw.strip()
    while text.startswith("[") and text.endswith("]"):
        text = text[1:-1].strip()
    if text.lower().startswith("context:"):
        text = text[len("context:") :].strip()
    return text


def _contact_line(contact: TailoredCVContact) -> str:
    parts = [contact.email, contact.phone]
    parts += [f'<a href="{_esc(link.url)}">{_esc(link.label)}</a>' for link in contact.links]
    return " &nbsp;|&nbsp; ".join(_esc(p) if not p.startswith("<a ") else p for p in parts if p)


def render_pdf(tailored: TailoredCV) -> bytes:
    """Render ``tailored`` as a polished, single-column A4 PDF."""

    styles = _styles()
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=16 * mm, bottomMargin=16 * mm,
        leftMargin=18 * mm, rightMargin=18 * mm,
        title="Tailored CV",
    )

    flow: list[object] = []
    contact = tailored.contact
    if contact is not None:
        flow.append(Paragraph(_esc(contact.name).upper(), styles["name"]))
        line = _contact_line(contact)
        if line:
            flow.append(Paragraph(line, styles["contact"]))
        if contact.location:
            flow.append(Paragraph(_esc(contact.location), styles["contact"]))
        flow.append(Spacer(1, 4))

    for section in tailored.sections:
        flow.extend(_section_flowables(section, styles))

    document.build(flow)
    return buffer.getvalue()


def _section_flowables(
    section: TailoredCVSection, styles: dict[str, ParagraphStyle]
) -> list[object]:
    """Heading + underline rule, then the section's bullets or entries.

    The heading, its rule, and the first content flowable are kept together so a
    heading is never orphaned at the bottom of a page; the rest flows freely.
    """

    content: list[object] = []
    if section.bullets:
        content.append(
            ListFlowable(
                [
                    ListItem(Paragraph(_esc(b), styles["bullet"]), leftIndent=10)
                    for b in section.bullets
                ],
                bulletType="bullet",
                start="•",
                leftIndent=12,
                bulletColor=_MUTED,
                spaceBefore=2,
                spaceAfter=4,
            )
        )
    for entry in section.entries:
        content.extend(_entry_flowables(entry, styles))

    head: list[object] = [
        Paragraph(_esc(section.heading).upper(), styles["heading"]),
        HRFlowable(
            width="100%", thickness=0.75, color=_RULE, spaceBefore=1, spaceAfter=6
        ),
    ]
    if content:
        head.append(content[0])
    return [KeepTogether(head), *content[1:]]


def _entry_flowables(
    entry: TailoredCVEntry, styles: dict[str, ParagraphStyle]
) -> list[object]:
    title = _esc(entry.title)
    if entry.organization:
        title = f"{title} <font color='#334155'>&mdash; {_esc(entry.organization)}</font>"
    left = Paragraph(f"<b>{title}</b>", styles["entry_title"])
    right = Paragraph(_esc(entry.date_range or ""), styles["entry_date"])
    header = Table([[left, right]], colWidths=[118 * mm, 46 * mm])
    header.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    out: list[object] = [header]
    if entry.context:
        out.append(
            Paragraph(f"[Context: {_esc(_context_text(entry.context))}]", styles["context"])
        )
    if entry.bullets:
        out.append(ListFlowable(
            [ListItem(Paragraph(_esc(b), styles["bullet"]), leftIndent=10)
             for b in entry.bullets],
            bulletType="bullet", start="•", leftIndent=12, bulletColor=_MUTED,
            spaceBefore=2, spaceAfter=4,
        ))
    else:
        out.append(Spacer(1, 3))
    return out


def render_docx(tailored: TailoredCV) -> bytes:
    """Render ``tailored`` as a clean DOCX document."""

    document = Document()
    contact = tailored.contact
    if contact is not None:
        h = document.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(contact.name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        meta_bits = [b for b in (contact.email, contact.phone) if b]
        meta_bits += [f"{link.label}: {link.url}" for link in contact.links]
        if contact.location:
            meta_bits.append(contact.location)
        if meta_bits:
            c = document.add_paragraph(" | ".join(meta_bits))
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in tailored.sections:
        document.add_heading(section.heading, level=2)
        for bullet in section.bullets:
            document.add_paragraph(bullet, style="List Bullet")
        for entry in section.entries:
            head = entry.title
            if entry.organization:
                head = f"{head} — {entry.organization}"
            if entry.date_range:
                head = f"{head}  ({entry.date_range})"
            p = document.add_paragraph()
            p.add_run(head).bold = True
            if entry.context:
                context_para = document.add_paragraph()
                context_para.add_run(
                    f"[Context: {_context_text(entry.context)}]"
                ).italic = True
            for bullet in entry.bullets:
                document.add_paragraph(bullet, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
